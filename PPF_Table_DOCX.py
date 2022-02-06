from docx import Document
import pandas as pd
from Functions import read_file
import matplotlib.pyplot as plt

plt.rcParams["font.family"] = "Century Gothic"
plt.rcParams["font.size"] = "14"
plt.rcParams['lines.linewidth'] = 3


def Calc_table(N_table, Folder, Name):
    doc = Document(f'{Folder}/{Name}')
    n_colums = len(doc.tables[N_table].rows[0].cells)
    n_rows = len(doc.tables[N_table].rows)
    n_round = 3
    for i in range(2, n_rows):
        Data = doc.tables[N_table].rows[i]
        min = float(Data.cells[1].text)
        typ = float(Data.cells[2].text)
        max = float(Data.cells[3].text)

        # sigma global data
        sigma_min = round((typ - min) / 3, n_round)
        sigma_max = round((max - typ) / 3, n_round)
        Data.cells[4].text = f'{sigma_min}'
        Data.cells[5].text = f'{sigma_max}'

        # mismatch sigma data
        sigma_mis = float(Data.cells[6].text)

        # total sigma data
        sigma_total_min = (sigma_min ** 2 + sigma_mis ** 2) ** 0.5
        sigma_total_max = (sigma_max ** 2 + sigma_mis ** 2) ** 0.5

        # min/max delta from tt (3sigma)
        min_delta = round(3 * sigma_total_min, n_round)
        max_delta = round(3 * sigma_total_max, n_round)
        Data.cells[7].text = f'-{min_delta}\n +{max_delta}'

        # min/max delta from tt (3sigma) in %
        min_percent = round(abs(typ - min_delta - typ) / (typ) * 100, n_round - 1)
        max_percent = round(abs(typ + max_delta - typ) / (typ) * 100, n_round - 1)
        if abs(min_percent) <= 100 or abs(max_percent) <= 100:
            Data.cells[8].text = f'-{abs(min_percent)}\n +{abs(max_percent)}'
        else:
            Data.cells[8].text = f''

    doc.save('C:/Users\ELECTRONIC\Desktop\PPF/PPF.docx')

def P1dB_calc(Home, Filter, Figure):
    P1dB_corners = ['SS', 'TT', 'FF']
    P1dB_min = []
    P1dB_max = []
    for i in P1dB_corners:
        P1dB_Data = read_file(f'{Home}/PPF_F/P1dB_{Filter}_{i}.csv')
        P1dB_min.append(min(P1dB_Data[:, 1]))
        P1dB_max.append(max(P1dB_Data[:, 1]))
        plt.figure(Figure)
        plt.title(Filter)
        plt.plot(P1dB_Data[:, 0], P1dB_Data[:, 1], label=f'{i}', linewidth='3')
    plt.legend()
    plt.xlabel('F, Гц')
    plt.ylabel('P1dB, дБм')
    plt.grid()
    print(f'{Filter} P1dB(min)={min(P1dB_min)}')
    print(f'{Filter} P1dB(max)={max(P1dB_max)}')

def HF_GD_Plot(Home, Filter, Figure):
    HF = read_file(f'{Home}/PPF_F/HF_{Filter}.csv')
    HF_M = read_file(f'{Home}/PPF_F/HF_{Filter}_Mirror.csv')
    GD = read_file(f'{Home}/PPF_F/GD_{Filter}.csv')
    #HF plot
    plt.figure(Figure)
    plt.title(Filter)
    plt.plot(HF[:, 0], HF[:, 1], color='tab:blue', linewidth='3')
    plt.plot(HF_M[:, 0], HF_M[:, 1], color='tab:blue', linewidth='3')
    plt.xlabel('F, Гц')
    plt.ylabel('H(F), дБ')
    plt.grid()
    #GD plot
    plt.figure(Figure+1)
    plt.title(Filter)
    plt.plot(GD[:, 0], GD[:, 1]*1e9, color='tab:blue', linewidth='3')
    plt.xlabel('F, Гц')
    plt.ylabel('GD(F), нс')
    plt.grid()


Home = 'C:/Users\ELECTRONIC/Desktop/PPF'
File = 'PPF.docx'

#P1dB_calc(Home, 'GPS', Figure=1)
#P1dB_calc(Home, 'GLO', Figure=2)

#HF_GD_Plot(Home, 'GPS', Figure=3)
#HF_GD_Plot(Home, 'GLO', Figure=5)

#Calc_table(N_table=1, Folder=Home, Name=File)
#Calc_table(N_table=3, Folder=Home, Name=File)

Step_1 = read_file(f'{Home}/PPF_F/Tran_GLO.csv')
Step_2 = read_file(f'{Home}/PPF_F/Tran_GPS.csv')
#HF plot
plt.figure(5)
plt.title('Tran')
plt.plot(Step_1[:, 0], Step_1[:, 1], color='tab:blue', label = 'GLO', linewidth='3')
plt.plot(Step_2[:, 0], Step_2[:, 1], color='tab:red', label = 'GPS', linewidth='3')
plt.xlabel('t, c')
plt.ylabel('h(t), В')
plt.legend()
plt.grid()




plt.show()







